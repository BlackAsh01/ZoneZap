"""
Convert PROJECT-REPORT.md to PROJECT-REPORT.docx using python-docx.
Run: py -3 md_to_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set table cell background (e.g. for header row)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_paragraph_with_format(doc, text):
    """Add paragraph, handling **bold** (asterisks not shown) and plain text."""
    if not text.strip():
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Split by **content** so odd-indexed segments are bold (non-greedy, so **a** **b** works)
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        part = part.replace('\n', ' ')
        if not part:
            continue
        run = p.add_run(part)
        if i % 2 == 1:
            run.bold = True
    return p

def add_code_paragraph(doc, code, style_name='Normal'):
    """Add a paragraph with monospace font for code."""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def parse_table(lines):
    """Parse markdown table lines into list of rows (list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith('|'):
            continue
        # Remove leading/trailing | and split by |
        cells = [c.strip() for c in line.strip('|').split('|')]
        # Skip separator row (|---|---|)
        if cells and re.match(r'^[\s\-:]+$', cells[0].replace('|', '')):
            continue
        rows.append(cells)
    return rows

def main():
    base = Path(__file__).resolve().parent
    md_path = base / 'PROJECT-REPORT.md'
    docx_path = base / 'PROJECT-REPORT.docx'

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    # Times New Roman and justified for body
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = 'Times New Roman'
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for heading_name in ('Heading 1', 'Heading 2', 'Heading 3'):
        if heading_name in doc.styles:
            doc.styles[heading_name].font.name = 'Times New Roman'

    lines = content.split('\n')
    i = 0
    in_code = False
    code_lang = ''
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line

        if in_code:
            if line.strip().startswith('```'):
                code_block = '\n'.join(code_lines)
                add_code_paragraph(doc, code_block)
                code_lines = []
                in_code = False
                i += 1
                continue
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith('```'):
            in_code = True
            code_lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            continue

        # Table: collect consecutive table rows
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                rows = parse_table(table_lines)
                if rows:
                    ncols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=ncols)
                    table.style = 'Table Grid'
                    for ri, row in enumerate(rows):
                        for ci, cell_text in enumerate(row):
                            if ci < ncols:
                                cell = table.rows[ri].cells[ci]
                                cell.text = cell_text
                                if ri == 0:
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.bold = True
                in_table = False
                table_lines = []

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('_' * 60)
            i += 1
            continue

        # List item
        if line.strip().startswith('- ') and not line.strip().startswith('- **'):
            add_paragraph_with_format(doc, '\u2022 ' + line.strip()[2:])
            i += 1
            continue
        if line.strip().startswith('- **'):
            add_paragraph_with_format(doc, '\u2022 ' + line.strip()[2:])
            i += 1
            continue

        # Numbered list (e.g. "1. ")
        if re.match(r'^\s*\d+\.\s', line):
            add_paragraph_with_format(doc, line.strip())
            i += 1
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Image: ![alt](path)
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line.strip())
        if img_match:
            alt, path = img_match.group(1), img_match.group(2)
            full_path = base / path
            if full_path.exists():
                try:
                    doc.add_picture(str(full_path), width=Inches(5.5))
                    if alt.strip():
                        p = doc.add_paragraph(alt.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph(f"[Image: {path}]")
            else:
                doc.add_paragraph(f"[Image: {path}]")
            i += 1
            continue

        # Normal paragraph (may contain **bold**)
        add_paragraph_with_format(doc, line)
        i += 1
        continue

    # Flush any remaining code block
    if in_code and code_lines:
        add_code_paragraph(doc, '\n'.join(code_lines))
    # Flush any remaining table
    if in_table and table_lines:
        rows = parse_table(table_lines)
        if rows:
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = 'Table Grid'
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci < ncols:
                        table.rows[ri].cells[ci].text = cell_text
                        if ri == 0:
                            for r in table.rows[ri].cells[ci].paragraphs[0].runs:
                                r.bold = True

    doc.save(docx_path)
    print(f"Created: {docx_path}")

if __name__ == '__main__':
    main()
